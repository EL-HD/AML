"""
generar_docs.py  —  AML Intelligence Platform
Genera documentacion_tecnica.docx en el directorio de trabajo actual.

Uso:
    python3 generar_docs.py

Requisitos:
    pip install python-docx
"""

from pathlib import Path
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─────────────────────────────────────────────────────────────
# HELPERS DE ESTILO
# ─────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Aplica color de fondo a una celda de tabla."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    """Aplica bordes a una celda."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{side}")
        tag.set(qn("w:val"), kwargs.get("val", "single"))
        tag.set(qn("w:sz"), kwargs.get("sz", "6"))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), kwargs.get("color", "CCCCCC"))
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = RGBColor(0x1A, 0x27, 0x44) if level == 1 else RGBColor(0x1E, 0x3A, 0x5F)
    return p


def body(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(10.5)
    return p


def add_table(doc: Document, headers: list, rows: list, col_widths: list = None):
    """Crea una tabla estilizada."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    # Cabecera
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_bg(hdr_cells[i], "1A2744")
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9.5)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_border(hdr_cells[i], color="1A2744")

    # Filas
    for idx, row_data in enumerate(rows):
        row_cells = table.add_row().cells
        bg = "F5F7FA" if idx % 2 == 0 else "FFFFFF"
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
            set_cell_bg(row_cells[i], bg)
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(9)
            set_cell_border(row_cells[i], color="D0D5DD")

    # Ancho de columnas
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table


# ─────────────────────────────────────────────────────────────
# CONSTRUCCIÓN DEL DOCUMENTO
# ─────────────────────────────────────────────────────────────

def build_document(output_path: Path):
    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # ── PORTADA ──────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SOVEREIGN AML")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1A, 0x27, 0x44)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle.add_run("Analytical Intelligence Platform — Documentación Técnica")
    run2.font.size = Pt(14)
    run2.font.color.rgb = RGBColor(0x34, 0x5E, 0x8A)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Versión 3.0  ·  {date.today().strftime('%d de %B de %Y')}").font.size = Pt(10)

    author_p = doc.add_paragraph()
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_a = author_p.add_run("Ing. Hobéd Díaz — Msc. M.A.F.I")
    run_a.font.size = Pt(10)
    run_a.font.color.rgb = RGBColor(0x55, 0x5E, 0x6D)

    doc.add_page_break()

    # ── 1. INTRODUCCIÓN ──────────────────────────────────────
    heading(doc, "1. Introducción", 1)
    body(doc, (
        "SOVEREIGN AML es una solución integral de inteligencia analítica y cumplimiento "
        "normativo de última generación. A diferencia de las herramientas reactivas tradicionales, "
        "esta plataforma utiliza análisis de datos avanzado y modelos de grafos para transformar el "
        "monitoreo de riesgos en oportunidades estratégicas de negocio, integrando estándares "
        "internacionales como ISO 31000 y el enfoque basado en riesgo (RBA) del GAFI."
    ))
    body(doc, (
        "La plataforma está dirigida a oficiales de cumplimiento, analistas AML y supervisores "
        "operativos que requieren una herramienta de alta fidelidad para el procesamiento de "
        "datos transaccionales complejos y la visualización de redes de lavado de activos."
    ))

    # ── 2. REQUISITOS ────────────────────────────────────────
    heading(doc, "2. Requisitos del Sistema", 1)
    heading(doc, "2.1 Software", 2)

    add_table(doc,
        headers=["Componente", "Versión mínima", "Descripción"],
        rows=[
            ["Python",         "3.9+",    "Lenguaje base del sistema"],
            ["Streamlit",      "1.30+",   "Framework de interfaz web premium"],
            ["pandas",         "2.0+",    "Procesamiento de datos vectoriales"],
            ["plotly",         "5.0+",    "Visualizaciones interactivas de alta densidad"],
            ["networkx",       "3.0+",    "Motor de análisis de grafos y redes"],
            ["numpy",          "1.24+",   "Cálculos numéricos y estadísticos"],
            ["openpyxl",       "3.1+",    "Lectura de archivos Excel (.xlsx)"],
            ["python-docx",    "1.0+",    "Generación de documentación (.docx)"],
            ["requests",       "2.28+",   "Integración con API de Autenticación y IA"],
        ],
        col_widths=[4.0, 3.5, 9.0]
    )

    heading(doc, "2.2 Archivo de Entrada", 2)
    body(doc, "La plataforma requiere un archivo Excel (.xlsx) con la siguiente estructura extendida:")
    add_table(doc,
        headers=["Columna", "Tipo", "Descripción"],
        rows=[
            ["Fecha",    "Fecha",       "Fecha de la transacción (YYYY-MM-DD)"],
            ["Cliente",  "Texto",       "Identificador único del cliente origen"],
            ["Monto",    "Numérico",    "Monto de la transacción (Q)"],
            ["Perfil",   "Numérico",    "Límite esperado de actividad (Q)"],
            ["EsPEP",    "Booleano",    "Persona Expuesta Políticamente"],
            ["EsCPE",    "Booleano",    "Contratista o Proveedor del Estado"],
            ["Ubicacion","Texto",       "Departamento o municipio de origen"],
            ["Cliente_Destino", "Texto", "Requerido para el análisis de Red Transaccional"],
        ],
        col_widths=[3.5, 3.0, 10.0]
    )

    # ── 3. INSTALACIÓN ───────────────────────────────────────
    heading(doc, "3. Instalación y Configuración", 1)
    heading(doc, "3.1 Instalación de Dependencias", 2)
    body(doc, "Ejecutar el siguiente comando para preparar el entorno:")

    code_p = doc.add_paragraph()
    code_run = code_p.add_run("pip install streamlit pandas plotly networkx numpy openpyxl python-docx requests")
    code_run.font.name = "Courier New"
    code_run.font.size = Pt(9.5)
    code_run.font.color.rgb = RGBColor(0x0D, 0x6E, 0xFD)

    heading(doc, "3.2 Ejecución del Ecosistema", 2)
    body(doc, "La plataforma requiere la API de Autenticación activa. En terminales separadas:")
    
    body(doc, "1. Iniciar Auth API:")
    code_p_api = doc.add_paragraph()
    run_api = code_p_api.add_run("python3 auth_api.py")
    run_api.font.name = "Courier New"
    run_api.font.size = Pt(9.5)
    run_api.font.color.rgb = RGBColor(0x0D, 0x6E, 0xFD)

    body(doc, "2. Iniciar SOVEREIGN UI:")
    code_p2 = doc.add_paragraph()
    run_c2 = code_p2.add_run("streamlit run app2.py")
    run_c2.font.name = "Courier New"
    run_c2.font.size = Pt(9.5)
    run_c2.font.color.rgb = RGBColor(0x0D, 0x6E, 0xFD)

    # ── 4. ARQUITECTURA TÉCNICA ──────────────────────────────
    heading(doc, "4. Arquitectura de Inteligencia (IMPERATOR ENGINE)", 1)
    body(doc, (
        "SOVEREIGN opera bajo el motor IMPERATOR, que segmenta el análisis en 4 dimensiones críticas:"
    ))

    capas = [
        ("S_T (Transaccional)",   "Detección de patrones mediante reglas de negocio (Smurfing, Picos, Acumulados)."),
        ("S_C (Contextual)",      "Evaluación del perfil del cliente: PEP, CPE y zonas geográficas de riesgo."),
        ("S_B (Conductual)",      "Análisis de desviaciones estadísticas respecto al comportamiento histórico."),
        ("S_N (Red)",             "Módulo de grafos que identifica cuentas puente, ciclos y estratificación."),
    ]
    add_table(doc,
        headers=["Dimensión", "Descripción"],
        rows=capas,
        col_widths=[5.0, 11.5]
    )

    # ── 5. MOTOR DE REGLAS AML ───────────────────────────────
    heading(doc, "5. Configuración de Detección", 1)
    body(doc, (
        "El motor de detección es dinámico y permite ajustar pesos y umbrales en tiempo real. "
        "Los parámetros por defecto del sistema son:"
    ))

    add_table(doc,
        headers=["Regla", "Variable", "Condición (default)", "Peso", "Activa"],
        rows=[
            ["Monto Crítico",   "Alerta_Absoluto", "Monto > Q20,000", "3",  "Sí"],
            ["Perfil Excedido", "Alerta_15",       "Monto > Perfil × 1.15", "1", "Sí"],
            ["Smurfing",        "Smurfing",         "Transacciones mismo día ≥ 5", "3", "Sí"],
            ["Geografía Riesgo","Ubicacion",        "Zonas fronterizas/sensibles", "2", "Sí"],
            ["Status Político", "PEP/CPE",          "Persona Expuesta / Contratista", "2", "Sí"],
        ],
        col_widths=[4.0, 4.0, 5.5, 2.0, 2.5]
    )

    # ── 6. CLASIFICACIÓN DE RIESGO ───────────────────────────
    heading(doc, "6. Niveles de Alerta SOVEREIGN", 1)
    add_table(doc,
        headers=["Nivel", "Score Requerido", "Acción de Mitigación"],
        rows=[
            ["Crítico", "≥ 8.0", "Bloqueo inmediato y Reporte de Operación Sospechosa (ROS)"],
            ["Alto",    "5.0 - 7.9", "Debida Diligencia Ampliada (EDD) obligatoria"],
            ["Medio",   "3.0 - 4.9", "Monitoreo reforzado y actualización de perfil"],
            ["Bajo",    "< 3.0", "Vigilancia rutinaria"],
        ],
        col_widths=[2.5, 6.5, 7.5]
    )

    # ── 7. VISTAS DEL SISTEMA ────────────────────────────────
    heading(doc, "7. Módulos de Inteligencia", 1)

    vistas = [
        ("📋 Resumen Ejecutivo",
         "KPIs centrales, distribución de riesgo y matriz de oportunidad de canal."),
        ("🚨 Casos de Alerta",
         "Centro de gestión de hallazgos con filtrado avanzado por score y tipo."),
        ("🌐 Red Transaccional",
         "Visualización de grafos para detectar estratificación y cuentas puente."),
        ("🛡️ Acciones de Mitigación",
         "Módulo de gestión de riesgos basado en el enfoque RBA del GAFI."),
        ("📍 Gestión de Ubicaciones",
         "Configuración dinámica de zonas de riesgo geográfico."),
        ("📊 Informes y Reportes",
         "Generación de reportes PDF de alta fidelidad para entes reguladores."),
        ("👤 Análisis por Cliente",
         "Perfil 360° con tendencias, mapas de calor y resumen generado por IA."),
        ("⚙️ Configuración",
         "Control total de umbrales, pesos y parámetros del motor IMPERATOR."),
    ]
    for nombre, desc in vistas:
        heading(doc, nombre, 2)
        body(doc, desc)

    # ── 8. NOTAS TÉCNICAS ───────────────────────────────────
    heading(doc, "8. Seguridad y Cumplimiento", 1)

    notas = [
        ("Control de Acceso", 
         "Autenticación robusta via Auth API con gestión de licencias corporativas."),
        ("OWASP Top 10",
         "Diseño orientado a la seguridad: prevención de inyecciones, control de acceso roto y fallos criptográficos."),
        ("IA Generativa",
         "Resúmenes analíticos automatizados mediante integración con modelos de lenguaje avanzados."),
        ("Portabilidad",
         "Estructura modular compatible con entornos locales y despliegues en la nube."),
    ]
    for titulo, desc in notas:
        heading(doc, titulo, 2)
        body(doc, desc)

    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_f = footer_p.add_run(f"SOVEREIGN AML Intelligence Platform v3.0  ·  Generado el {date.today().strftime('%d/%m/%Y')}")
    run_f.font.size = Pt(9)
    run_f.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # Guardar en directorio actual
    doc.save(str(output_path))
    print(f"✅ Documento generado: {output_path}")
    return output_path


if __name__ == "__main__":
    output = Path.cwd() / "documentacion_tecnica.docx"
    build_document(output)
