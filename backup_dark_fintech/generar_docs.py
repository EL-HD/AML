"""
generar_docs.py  —  AML Intelligence Platform
Genera documentacion_tecnica.docx en el directorio de trabajo actual.

Uso:
    python3 generar_docs.py

Requisitos:
    pip install python-docx
"""

from pathlib import Path
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─────────────────────────────────────────────────────────────
# HELPERS DE ESTILO
# ─────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Aplica color de fondo a una celda de tabla."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    """Aplica bordes a una celda."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{side}")
        tag.set(qn("w:val"), kwargs.get("val", "single"))
        tag.set(qn("w:sz"), kwargs.get("sz", "6"))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), kwargs.get("color", "CCCCCC"))
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = RGBColor(0x1A, 0x27, 0x44) if level == 1 else RGBColor(0x1E, 0x3A, 0x5F)
    return p


def body(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(10.5)
    return p


def add_table(doc: Document, headers: list, rows: list, col_widths: list = None):
    """Crea una tabla estilizada."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    # Cabecera
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_bg(hdr_cells[i], "1A2744")
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9.5)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_border(hdr_cells[i], color="1A2744")

    # Filas
    for idx, row_data in enumerate(rows):
        row_cells = table.add_row().cells
        bg = "F5F7FA" if idx % 2 == 0 else "FFFFFF"
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
            set_cell_bg(row_cells[i], bg)
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(9)
            set_cell_border(row_cells[i], color="D0D5DD")

    # Ancho de columnas
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table


# ─────────────────────────────────────────────────────────────
# CONSTRUCCIÓN DEL DOCUMENTO
# ─────────────────────────────────────────────────────────────

def build_document(output_path: Path):
    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # ── PORTADA ──────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AML INTELLIGENCE PLATFORM")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x1A, 0x27, 0x44)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle.add_run("Documentación Técnica de Ejecución y Uso")
    run2.font.size = Pt(14)
    run2.font.color.rgb = RGBColor(0x34, 0x5E, 0x8A)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Versión 2.5  ·  {date.today().strftime('%d de %B de %Y')}").font.size = Pt(10)

    author_p = doc.add_paragraph()
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_a = author_p.add_run("Ing. Hobéd Díaz — Msc. M.A.F.I")
    run_a.font.size = Pt(10)
    run_a.font.color.rgb = RGBColor(0x55, 0x5E, 0x6D)

    doc.add_page_break()

    # ── 1. INTRODUCCIÓN ──────────────────────────────────────
    heading(doc, "1. Introducción", 1)
    body(doc, (
        "AML Intelligence Platform es una solución integral de Business Intelligence (BI) y cumplimiento "
        "normativo. A diferencia de las herramientas reactivas tradicionales, esta plataforma utiliza "
        "análisis de datos avanzado para transformar el monitoreo de riesgos en oportunidades "
        "estratégicas de negocio, permitiendo identificar tendencias de mercado y proyecciones de "
        "crecimiento mientras se mantiene una postura de seguridad impecable."
    ))
    body(doc, (
        "La plataforma está dirigida a oficiales de cumplimiento, analistas AML y supervisores "
        "operativos de instituciones financieras que requieren una herramienta ágil para el "
        "procesamiento de datos transaccionales y la generación de reportes de riesgo."
    ))

    # ── 2. REQUISITOS ────────────────────────────────────────
    heading(doc, "2. Requisitos del Sistema", 1)
    heading(doc, "2.1 Software", 2)

    add_table(doc,
        headers=["Componente", "Versión mínima", "Descripción"],
        rows=[
            ["Python",         "3.9+",    "Lenguaje base del sistema"],
            ["Streamlit",      "1.30+",   "Framework de interfaz web"],
            ["pandas",         "2.0+",    "Procesamiento de datos tabulares"],
            ["plotly",         "5.0+",    "Visualizaciones interactivas con hover"],
            ["matplotlib",     "3.7+",    "Gráficos auxiliares (compatibilidad)"],
            ["numpy",          "1.24+",   "Cálculos numéricos y estadísticos"],
            ["openpyxl",       "3.1+",    "Lectura de archivos Excel (.xlsx)"],
            ["python-docx",    "1.0+",    "Generación de documentación (.docx)"],
            ["requests",       "2.28+",   "Integración con API de IA (Claude)"],
        ],
        col_widths=[4.0, 3.5, 9.0]
    )

    heading(doc, "2.2 Archivo de Entrada", 2)
    body(doc, "La plataforma requiere un archivo Excel (.xlsx) con las siguientes columnas obligatorias:")
    add_table(doc,
        headers=["Columna", "Tipo", "Descripción"],
        rows=[
            ["Cliente",  "Texto / ID",  "Identificador único del cliente"],
            ["Monto",    "Numérico",    "Monto de la transacción en quetzales (Q)"],
            ["Perfil",   "Numérico",    "Límite esperado de actividad del cliente (Q)"],
            ["Fecha",    "Fecha",       "Fecha de la transacción (formato YYYY-MM-DD o compatible)"],
        ],
        col_widths=[3.5, 3.0, 10.0]
    )

    # ── 3. INSTALACIÓN ───────────────────────────────────────
    heading(doc, "3. Instalación y Configuración", 1)
    heading(doc, "3.1 Instalación de Dependencias", 2)
    body(doc, "Ejecutar el siguiente comando desde el directorio del proyecto:")

    code_p = doc.add_paragraph()
    code_run = code_p.add_run("pip install streamlit pandas plotly matplotlib numpy openpyxl python-docx requests")
    code_run.font.name = "Courier New"
    code_run.font.size = Pt(9.5)
    code_run.font.color.rgb = RGBColor(0x0D, 0x6E, 0xFD)

    heading(doc, "3.2 Ejecución de la Plataforma", 2)
    body(doc, "Desde el directorio donde se encuentra app2.py, ejecutar:")

    code_p2 = doc.add_paragraph()
    run_c2 = code_p2.add_run("streamlit run app2.py")
    run_c2.font.name = "Courier New"
    run_c2.font.size = Pt(9.5)
    run_c2.font.color.rgb = RGBColor(0x0D, 0x6E, 0xFD)

    body(doc, (
        "Streamlit abrirá automáticamente el navegador en http://localhost:8501. "
        "Si el navegador no se abre, ingresar esa URL manualmente. "
        "Para detener el servidor, presionar Ctrl+C en la terminal."
    ))

    heading(doc, "3.3 Generación de la Documentación", 2)
    body(doc, "Para regenerar este documento, ejecutar:")
    code_p3 = doc.add_paragraph()
    run_c3 = code_p3.add_run("python3 generar_docs.py")
    run_c3.font.name = "Courier New"
    run_c3.font.size = Pt(9.5)
    run_c3.font.color.rgb = RGBColor(0x0D, 0x6E, 0xFD)
    body(doc, "El archivo documentacion_tecnica.docx se generará en el directorio actual de trabajo.")

    # ── 4. ARQUITECTURA TÉCNICA ──────────────────────────────
    heading(doc, "4. Arquitectura Técnica", 1)
    body(doc, (
        "La plataforma sigue una arquitectura de script único (single-file app) con las siguientes capas lógicas:"
    ))

    capas = [
        ("Carga de datos",        "Lectura del archivo Excel mediante pandas y validación de columnas requeridas."),
        ("Motor de reglas",       "Aplicación secuencial de 6 reglas AML parametrizables usando operaciones vectoriales de pandas."),
        ("Scoring",               "Suma ponderada de reglas activas por transacción, con peso configurable por regla (0–12 pts)."),
        ("Clasificación",         "Asignación del nivel de riesgo (Crítico / Alto / Medio / Bajo) por cliente basada en Score_Max y Total_Mensual."),
        ("Inteligencia (BI)",     "Capa de análisis estratégico que cruza riesgo vs. volumen para identificar canales de alto rendimiento."),
        ("Visualización",         "Gráficos interactivos con Plotly (hover, zoom, descarga) y cuadros de interpretación de negocio."),
        ("Integración IA",        "Generación de un resumen analítico textual por cliente mediante API de Claude (Anthropic)."),
    ]
    add_table(doc,
        headers=["Capa", "Descripción"],
        rows=capas,
        col_widths=[5.0, 11.5]
    )

    # ── 5. MOTOR DE REGLAS AML ───────────────────────────────
    heading(doc, "5. Motor de Reglas AML", 1)
    body(doc, (
        "El sistema evalúa 6 reglas independientes sobre cada transacción. Los parámetros son "
        "completamente configurables desde la vista '⚙️ Configuración de Reglas'. "
        "A continuación se describen las reglas con sus valores por defecto:"
    ))

    add_table(doc,
        headers=["Regla", "Variable", "Condición (default)", "Peso default", "Activa por default"],
        rows=[
            ["Monto Alto Absoluto", "Alerta_Absoluto", "Monto > Q20,000", "3",  "Sí"],
            ["Acumulado Mensual",   "Alerta_Acumulado","Total_Mensual > Perfil × 2.0", "2", "Sí"],
            ["Exceso sobre Perfil", "Alerta_15",       "Monto > Perfil × 1.15", "1", "Sí"],
            ["Frecuencia Alta",     "Alerta_Frecuencia","Transacciones en período > 5", "1", "Sí"],
            ["Smurfing",            "Smurfing",         "Transacciones en mismo día ≥ 5", "3", "Sí"],
            ["Pico Anómalo",        "Pico",             "Monto > Media + 2×Desv. Estándar", "2", "Sí"],
        ],
        col_widths=[4.0, 4.0, 5.5, 2.0, 2.5]
    )

    doc.add_paragraph()
    body(doc, "Fórmula de Score por transacción:")
    code_score = doc.add_paragraph()
    r = code_score.add_run(
        "Score = (Alerta_Absoluto × peso_absoluto) + (Alerta_Acumulado × peso_acumulado)\n"
        "      + (Alerta_15 × peso_perfil)  + (Alerta_Frecuencia × peso_frecuencia)\n"
        "      + (Smurfing × peso_smurfing) + (Pico × peso_pico)"
    )
    r.font.name = "Courier New"
    r.font.size = Pt(9)

    # ── 6. CLASIFICACIÓN DE RIESGO ───────────────────────────
    heading(doc, "6. Clasificación de Riesgo por Cliente", 1)
    body(doc, (
        "El nivel de riesgo final se asigna según el Score_Max (máximo score alcanzado "
        "por cualquier transacción del cliente) y el Total_Mensual:"
    ))

    add_table(doc,
        headers=["Nivel", "Condición (default)", "Acción recomendada"],
        rows=[
            ["Crítico", "Score_Max ≥ 8 O Total_Mensual > Q30,000", "Revisión inmediata, posible reporte a UIF"],
            ["Alto",    "Score_Max ≥ 5",                           "Investigación obligatoria en 5 días hábiles"],
            ["Medio",   "Score_Max ≥ 3",                           "Monitoreo continuo, revisión mensual"],
            ["Bajo",    "Score_Max < 3",                           "Vigilancia rutinaria"],
        ],
        col_widths=[2.5, 6.5, 7.5]
    )

    # ── 7. VISTAS DEL SISTEMA ────────────────────────────────
    heading(doc, "7. Vistas del Sistema", 1)

    vistas = [
        ("📋 Resumen Ejecutivo",
         "Vista de alto nivel con KPIs principales y un Panel de Insights Estratégicos. Incluye gráficos "
         "de distribución de riesgo, barras de alertas por tipo y la Matriz de Oportunidad de Canal."),
        ("🎯 Inteligencia Estratégica",
         "Módulo que visualiza la Matriz de Oportunidad (Burbujas: Volumen vs Clientes) y la Densidad de "
         "Riesgo Agrupada para identificar Canales de Alto Rendimiento."),
        ("🚨 Casos de Alerta",
         "Tabla filtrable de todos los clientes con su Score_Max, Total_Mensual, número de transacciones "
         "y nivel de riesgo. Soporta filtros por nivel de riesgo (multiselect) y score mínimo (slider)."),
        ("📄 Transacciones",
         "Registro detallado de todas las transacciones con sus columnas de alerta booleanas y score "
         "individual. Permite filtrar por cliente y mostrar solo transacciones con alertas activas."),
        ("👤 Análisis por Cliente",
         "Perfil individual del cliente: KPIs, resumen generado por IA (Claude API), gráfico de tendencia "
         "de montos vs. perfil, mapa de picos anómalos estadísticos y frecuencia diaria de operaciones."),
        ("🗂️ Matrices de Riesgo",
         "Matriz de todos los clientes ordenada por score descendente y matriz de tipos de alerta con "
         "descripción, peso y cantidad de activaciones. Incluye gráfico de contribución al score por regla."),
        ("⚙️ Configuración de Reglas",
         "Panel de control completo con 4 pestañas: Reglas de Detección (umbrales individuales con "
         "toggle on/off), Pesos del Score, Clasificación de Riesgo y Resumen con botón de aplicar. "
         "Todos los cambios son persistentes en la session_state de Streamlit."),
    ]
    for nombre, desc in vistas:
        heading(doc, nombre, 2)
        body(doc, desc)

    # ── 8. GUÍA DE USO ───────────────────────────────────────
    heading(doc, "8. Guía de Uso Paso a Paso", 1)

    pasos = [
        ("Paso 1 — Preparar el archivo Excel",
         "Verificar que el archivo .xlsx tenga las columnas: Cliente, Monto, Perfil, Fecha. "
         "Los nombres son case-sensitive. Las fechas deben estar en formato compatible con pandas."),
        ("Paso 2 — Subir el archivo",
         "En la sección '📁 Cargar Datos', hacer clic en el área de carga y seleccionar el archivo. "
         "El sistema procesará automáticamente todas las reglas AML."),
        ("Paso 3 — Revisar el Resumen Ejecutivo",
         "La primera vista muestra los indicadores clave. Pasar el cursor sobre los gráficos "
         "para ver los valores exactos mediante tooltips interactivos."),
        ("Paso 4 — Analizar Casos Críticos",
         "Ir a '🚨 Casos de Alerta', filtrar por nivel 'Crítico' y exportar si es necesario."),
        ("Paso 5 — Investigar Clientes",
         "En '👤 Análisis por Cliente', seleccionar el cliente y usar el botón '🤖 Generar Resumen con IA' "
         "(requiere conexión a internet y API key configurada en el código)."),
        ("Paso 6 — Ajustar Parámetros",
         "Si los umbrales no se adaptan al perfil de la institución, ir a '⚙️ Configuración de Reglas' "
         "y ajustar. Luego volver a cargar el archivo para reprocesar con nuevos parámetros."),
    ]
    for titulo, desc in pasos:
        heading(doc, titulo, 2)
        body(doc, desc)

    # ── 9. GLOSARIO TÉCNICO ──────────────────────────────────
    heading(doc, "9. Glosario Técnico AML", 1)
    add_table(doc,
        headers=["Término", "Definición"],
        rows=[
            ["AML (Anti-Money Laundering)", "Conjunto de leyes, regulaciones y procedimientos para prevenir el lavado de activos."],
            ["Score de Riesgo",  "Puntuación compuesta (0–12) que refleja el nivel de actividad sospechosa de un cliente."],
            ["Smurfing",         "Técnica de fragmentar grandes sumas en múltiples transacciones pequeñas para evadir controles."],
            ["Pico Anómalo",     "Transacción cuyo monto supera la media histórica del cliente más 2 desviaciones estándar."],
            ["Perfil de Riesgo", "Nivel de actividad transaccional esperado para un cliente según su tipo de negocio."],
            ["Score_Max",        "Máximo score obtenido por cualquier transacción individual de un cliente en el período."],
            ["UIF",              "Unidad de Información Financiera — entidad receptora de reportes de operaciones sospechosas."],
            ["Session State",    "Mecanismo de Streamlit para mantener variables persistentes entre reruns de la aplicación."],
            ["Hover Tooltip",    "Información emergente que aparece al posicionar el cursor sobre un punto del gráfico Plotly."],
        ],
        col_widths=[5.5, 11.0]
    )

    # ── 10. NOTAS TÉCNICAS ───────────────────────────────────
    heading(doc, "10. Notas Técnicas Adicionales", 1)

    notas = [
        ("Gráficos interactivos", 
         "Todos los gráficos utilizan Plotly. Cada punto, barra o segmento muestra un tooltip al "
         "posicionar el cursor. Los gráficos soportan zoom (arrastrar), paneo (shift+arrastrar) "
         "y descarga como PNG mediante el botón de cámara en la esquina superior derecha."),
        ("Portabilidad",
         "La plataforma es completamente portable. Solo requiere Python 3.9+ y las dependencias "
         "listadas en la Sección 2. No requiere instalación de base de datos ni servicios externos "
         "(excepto la integración opcional con la API de IA)."),
        ("Integración con IA",
         "El módulo de resumen IA usa la API de Anthropic (Claude). Requiere que la API key esté "
         "configurada directamente en el código. Si no hay conexión o key válida, el resto de la "
         "plataforma funciona con normalidad."),
        ("Persistencia de datos",
         "La plataforma no almacena datos entre sesiones. Cada vez que se sube un nuevo archivo, "
         "se reprocesa todo desde cero. Los cambios de configuración se mantienen durante la sesión "
         "activa mediante session_state."),
    ]
    for titulo, desc in notas:
        heading(doc, titulo, 2)
        body(doc, desc)

    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_f = footer_p.add_run(f"AML Intelligence Platform v2.0  ·  Generado el {date.today().strftime('%d/%m/%Y')}")
    run_f.font.size = Pt(9)
    run_f.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # Guardar en directorio actual
    doc.save(str(output_path))
    print(f"✅ Documento generado: {output_path}")
    return output_path


if __name__ == "__main__":
    output = Path.cwd() / "documentacion_tecnica.docx"
    build_document(output)
